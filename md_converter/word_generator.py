"""
Word Generator используя python-docx
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from bs4 import BeautifulSoup
from typing import Optional
import re


class WordGenerator:
    """Генератор Word документов"""

    def __init__(self, style: str = 'default'):
        """
        Args:
            style: Стиль документа
        """
        self.style = style
        self.doc = Document()
        self._setup_styles()

    def generate(self, html_content: str, output_path: str, metadata: Optional[dict] = None):
        """
        Генерация Word документа из HTML

        Args:
            html_content: HTML контент
            output_path: Путь для сохранения
            metadata: Метаданные
        """
        # Добавление метаданных
        if metadata:
            self._add_metadata(metadata)

        # Парсинг HTML
        soup = BeautifulSoup(html_content, 'html.parser')

        # Обработка элементов
        self._process_elements(soup)

        # Сохранение документа
        self.doc.save(output_path)

    def _setup_styles(self):
        """Настройка стилей документа"""
        # Настройка базового стиля
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(11)

        # Настройка стилей заголовков
        for i in range(1, 7):
            heading_style = self.doc.styles[f'Heading {i}']
            heading_style.font.bold = True
            heading_style.font.color.rgb = RGBColor(44, 62, 80)

    def _add_metadata(self, metadata: dict):
        """Добавление метаданных документа"""
        core_properties = self.doc.core_properties

        if 'title' in metadata:
            core_properties.title = metadata['title']
            # Добавление заголовка в документ
            heading = self.doc.add_heading(metadata['title'], level=0)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if 'author' in metadata:
            core_properties.author = metadata['author']

        if 'date' in metadata:
            # Можно добавить как параграф
            self.doc.add_paragraph(f"Дата: {metadata['date']}")

        # Разделитель
        self.doc.add_paragraph()

    def _process_elements(self, soup):
        """Обработка HTML элементов"""
        for element in soup.find_all(recursive=False):
            self._process_element(element)

    def _process_element(self, element):
        """Обработка одного элемента"""
        tag_name = element.name

        if tag_name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            self._add_heading(element)
        elif tag_name == 'p':
            self._add_paragraph(element)
        elif tag_name == 'pre':
            self._add_code_block(element)
        elif tag_name == 'table':
            self._add_table(element)
        elif tag_name == 'ul':
            self._add_list(element, ordered=False)
        elif tag_name == 'ol':
            self._add_list(element, ordered=True)
        elif tag_name == 'blockquote':
            self._add_blockquote(element)
        elif tag_name == 'img':
            self._add_image(element)

    def _add_heading(self, element):
        """Добавление заголовка"""
        level = int(element.name[1])  # h1 -> 1, h2 -> 2, etc.
        text = element.get_text(strip=True)
        self.doc.add_heading(text, level=level)

    def _add_paragraph(self, element):
        """Добавление параграфа"""
        text = element.get_text(strip=True)
        if text:
            p = self.doc.add_paragraph(text)

    def _add_code_block(self, element):
        """Добавление блока кода"""
        code_text = element.get_text()

        # Создание параграфа с моноширинным шрифтом
        p = self.doc.add_paragraph()
        run = p.add_run(code_text)
        run.font.name = 'Courier New'
        run.font.size = Pt(9)

        # Фон (через XML)
        shading_elm = run._element.get_or_add_rPr()

        # Установка стиля как код
        p.style = 'No Spacing'

    def _add_table(self, element):
        """Добавление таблицы"""
        rows = element.find_all('tr')

        if not rows:
            return

        # Определение количества колонок
        first_row = rows[0]
        cols_count = len(first_row.find_all(['th', 'td']))

        # Создание таблицы
        table = self.doc.add_table(rows=0, cols=cols_count)
        table.style = 'Light Grid Accent 1'

        # Обработка строк
        for row_idx, row in enumerate(rows):
            cells = row.find_all(['th', 'td'])
            doc_row = table.add_row()

            for col_idx, cell in enumerate(cells):
                cell_text = cell.get_text(strip=True)
                doc_cell = doc_row.cells[col_idx]
                doc_cell.text = cell_text

                # Заголовки таблицы (жирный шрифт)
                if row_idx == 0 or cell.name == 'th':
                    for paragraph in doc_cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True

    def _add_list(self, element, ordered=False):
        """Добавление списка"""
        items = element.find_all('li', recursive=False)

        style = 'List Number' if ordered else 'List Bullet'

        for item in items:
            text = item.get_text(strip=True)
            self.doc.add_paragraph(text, style=style)

    def _add_blockquote(self, element):
        """Добавление цитаты"""
        text = element.get_text(strip=True)
        p = self.doc.add_paragraph(text)
        p.style = 'Intense Quote'

    def _add_image(self, element):
        """Добавление изображения"""
        src = element.get('src', '')

        if src and src.startswith('/') or src.startswith('./'):
            try:
                self.doc.add_picture(src, width=Inches(5))
            except Exception as e:
                # Если не удалось загрузить изображение, добавляем заглушку
                self.doc.add_paragraph(f'[Изображение: {src}]')
